from docx import Document


def extract_docx(docx_path: str, out_path: str) -> None:
    doc = Document(docx_path)
    lines: list[str] = []

    # Párrafos
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)

    # Tablas (mejor esfuerzo)
    for ti, tbl in enumerate(doc.tables, start=1):
        lines.append(f"[TABLA {ti}]")
        for row in tbl.rows:
            cells = []
            for c in row.cells:
                ct = (c.text or "").strip()
                if ct:
                    cells.append(ct)
            if cells:
                lines.append(" | ".join(cells))

    # Compactar blancos extra
    compacted: list[str] = []
    blank = 0
    for ln in lines:
        if ln.strip() == "":
            blank += 1
        else:
            blank = 0
        if blank <= 1:
            compacted.append(ln)

    out_text = "\n".join(compacted).strip() + "\n"
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(out_text)


if __name__ == "__main__":
    extract_docx(
        r"c:\Users\aldyj\Downloads\GLAB-S11-JFARFAN-2026-01.docx",
        "doc_extracted_full.txt",
    )
    print("OK -> doc_extracted_full.txt")

